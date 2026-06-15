# FILE: tests/test_documents.py
# Purpose: Unit tests for app/documents — xlsx/csv/docx/md ⇄ snapshot round
#          trips (real files in tmp), editor action queue, atomic save + .bak.
# Called-by: pytest
# Depends-on: app.documents.*
# Last-renovated: 2026-06-12
from __future__ import annotations

import asyncio

import pytest

from app.documents import (
    csv_md_convert,
    docx_convert,
    editor_actions,
    storage,
    xlsx_convert,
)


# ── xlsx round trip ────────────────────────────────────────────────────────

def _make_xlsx(path):
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "Data"
    ws["A1"] = "label"
    ws["B1"] = 41.5
    ws["B2"] = 8.5
    ws["B3"] = "=B1+B2"
    ws["C1"] = 0.25
    ws["C1"].number_format = "0.00%"
    ws.merge_cells("A5:B6")
    wb.save(path)


def test_xlsx_round_trip(tmp_path):
    source = tmp_path / "book.xlsx"
    _make_xlsx(str(source))

    snapshot = xlsx_convert.xlsx_to_snapshot(str(source))
    sheet = snapshot["sheets"][snapshot["sheetOrder"][0]]
    cells = sheet["cellData"]
    assert cells["0"]["0"]["v"] == "label"
    assert cells["0"]["1"]["v"] == 41.5 and cells["0"]["1"]["t"] == 2
    assert cells["2"]["1"]["f"] == "=B1+B2"
    # number format pattern survives via the styles table
    style_id = cells["0"]["2"]["s"]
    assert snapshot["styles"][style_id]["n"]["pattern"] == "0.00%"
    assert {"startRow": 4, "endRow": 5, "startColumn": 0, "endColumn": 1} \
        in sheet["mergeData"]

    out = tmp_path / "book_out.xlsx"
    xlsx_convert.snapshot_to_xlsx(snapshot, str(out))

    from openpyxl import load_workbook
    reloaded = load_workbook(str(out))
    ws = reloaded.active
    assert ws["A1"].value == "label"
    assert ws["B3"].value == "=B1+B2"
    assert ws["C1"].number_format == "0.00%"
    assert any(str(r) == "A5:B6" for r in ws.merged_cells.ranges)


# ── csv round trip ─────────────────────────────────────────────────────────

def test_csv_round_trip(tmp_path):
    source = tmp_path / "data.csv"
    source.write_text("name,count\nwidgets,41\nsprockets,8.5\n", encoding="utf-8")
    snapshot = csv_md_convert.csv_to_snapshot(str(source))
    cells = snapshot["sheets"]["sheet-01"]["cellData"]
    assert cells["0"]["0"]["v"] == "name"
    assert cells["1"]["1"]["v"] == 41 and cells["1"]["1"]["t"] == 2
    assert cells["2"]["1"]["v"] == 8.5

    out = tmp_path / "data_out.csv"
    csv_md_convert.snapshot_to_csv(snapshot, str(out))
    assert out.read_text(encoding="utf-8").strip() == \
        "name,count\nwidgets,41\nsprockets,8.5"


# ── docx round trip ────────────────────────────────────────────────────────

def _make_docx(path):
    import docx
    document = docx.Document()
    document.add_heading("Quarterly Report", level=1)
    p = document.add_paragraph()
    p.add_run("Revenue was ")
    p.add_run("strong").bold = True
    p.add_run(" this ")
    p.add_run("quarter").italic = True
    p.add_run(".")
    document.add_paragraph("First point", style="List Bullet")
    document.add_paragraph("Second point", style="List Bullet")
    document.save(path)


def test_docx_round_trip(tmp_path):
    source = tmp_path / "report.docx"
    _make_docx(str(source))

    snapshot = docx_convert.docx_to_snapshot(str(source))
    stream = snapshot["body"]["dataStream"]
    assert "Quarterly Report" in stream
    assert "• First point" in stream          # bullets surface visibly
    heading_runs = [r for r in snapshot["body"]["textRuns"]
                    if (r["ts"].get("fs") or 0) >= docx_convert.H1_SIZE]
    assert heading_runs, "heading must carry the H1 size encoding"
    bold_inline = [r for r in snapshot["body"]["textRuns"]
                   if r["ts"].get("bl") == 1 and "fs" not in r["ts"]]
    assert bold_inline, "inline bold run expected"

    out = tmp_path / "report_out.docx"
    docx_convert.snapshot_to_docx(snapshot, str(out))

    import docx
    reloaded = docx.Document(str(out))
    styles = [p.style.name for p in reloaded.paragraphs]
    texts = [p.text for p in reloaded.paragraphs]
    assert any(s.startswith("Heading 1") for s in styles)
    assert "First point" in texts and "Second point" in texts
    bullet_idx = texts.index("First point")
    assert reloaded.paragraphs[bullet_idx].style.name == "List Bullet"
    # bold survives
    body_para = next(p for p in reloaded.paragraphs if "Revenue" in p.text)
    assert any(r.bold for r in body_para.runs)


def test_complex_docx_opens_without_crashing(tmp_path):
    """Acceptance 5: graceful degradation, never a crash."""
    import docx
    source = tmp_path / "complex.docx"
    document = docx.Document()
    document.add_heading("Has a table", level=1)
    table = document.add_table(rows=2, cols=3)
    table.cell(0, 0).text = "inside a table"
    document.add_paragraph("after the table")
    document.save(str(source))

    snapshot = docx_convert.docx_to_snapshot(str(source))   # must not raise
    assert "after the table" in snapshot["body"]["dataStream"]


# ── markdown round trip ────────────────────────────────────────────────────

def test_md_round_trip(tmp_path):
    source = tmp_path / "notes.md"
    source.write_text(
        "# Title\nPlain line with **bold** and *lean* words.\n- alpha\n- beta\n",
        encoding="utf-8")
    snapshot = csv_md_convert.md_to_snapshot(str(source))
    stream = snapshot["body"]["dataStream"]
    assert "Title" in stream and "**" not in stream   # markers stripped to runs

    out = tmp_path / "notes_out.md"
    csv_md_convert.snapshot_to_md(snapshot, str(out))
    text = out.read_text(encoding="utf-8")
    assert text.splitlines()[0] == "# Title"
    assert "**bold**" in text and "*lean*" in text
    assert "- alpha" in text and "- beta" in text


# ── editor action queue ────────────────────────────────────────────────────

@pytest.mark.asyncio
async def test_execute_without_open_editor_is_honest():
    editor_actions.set_editor_state(False)
    out = await editor_actions.execute("doc_get_text", {})
    assert out["ok"] is False and "no document" in out["error"]


@pytest.mark.asyncio
async def test_action_round_trip_and_timeout():
    editor_actions.set_editor_state(True, path="x.docx", kind="doc", name="x.docx")
    try:
        async def pane():
            action = await editor_actions.next_action(wait_seconds=5)
            assert action["action_type"] == "sheet_set_range"
            editor_actions.post_result(action["action_id"],
                                       {"ok": True, "result": {"rows": 1}})

        pane_task = asyncio.create_task(pane())
        out = await editor_actions.execute("sheet_set_range",
                                           {"a1": "A10", "values": [["total"]]},
                                           timeout_seconds=5)
        await pane_task
        assert out["ok"] is True and out["result"]["rows"] == 1

        # nobody answering -> clean timeout, not a hang
        slow = await editor_actions.execute("doc_get_text", {}, timeout_seconds=0.2)
        assert slow["ok"] is False and "didn't answer" in slow["error"]
    finally:
        editor_actions.set_editor_state(False)


# ── storage discipline ─────────────────────────────────────────────────────

def test_atomic_write_and_bak_once(tmp_path):
    target = tmp_path / "doc.md"
    target.write_text("original", encoding="utf-8")

    created = storage.ensure_first_save_backup(str(target))
    assert created and storage.backup_path(str(target)).read_text(encoding="utf-8") == "original"

    storage.atomic_write_via(str(target),
                             lambda tmp: open(tmp, "w", encoding="utf-8").write("v2") and None)
    assert target.read_text(encoding="utf-8") == "v2"

    # second save: bak untouched
    created_again = storage.ensure_first_save_backup(str(target))
    assert created_again is False
    assert storage.backup_path(str(target)).read_text(encoding="utf-8") == "original"


def test_document_tools_register_cleanly():
    from app.tools import document_tools
    from app.tools.registry import list_tools
    document_tools.register_document_tools()
    names = {t["name"] for t in list_tools()}
    assert {"editor_status", "read_document", "edit_document_text",
            "read_sheet_range", "set_sheet_range", "list_sheet_names",
            "save_document"} <= names
