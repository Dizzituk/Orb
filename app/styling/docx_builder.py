# FILE: app/styling/docx_builder.py
# Purpose: Build styled Word documents from a structured content list.
# Called-by: app.debug.executors.styled_files, scripts.legal_case._rebuild_docx
# Depends-on: stdlib/third-party only
# Last-renovated: 2026-06-11
"""
Build styled Word documents from a structured content list.

Content is a list of "blocks":
  {"type": "heading", "level": 1, "text": "..."}
  {"type": "paragraph", "text": "..."}
  {"type": "list", "items": ["...", "..."], "ordered": False}
  {"type": "table", "headers": [...], "rows": [[...], [...]]}
  {"type": "rule"}            -- horizontal divider
  {"type": "spacer"}           -- blank line
  {"type": "code", "text": "...", "language": "python"}

Cover page is added automatically when theme.show_cover_page is True.
"""
from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional


def build_docx(
    output_path: str,
    title: str,
    content: List[Dict[str, Any]],
    theme: Dict[str, Any],
    subtitle: Optional[str] = None,
    author: Optional[str] = None,
) -> Dict[str, Any]:
    """Render a styled .docx file. Returns {path, size_bytes, blocks_rendered}."""
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # Page margins
    for section in doc.sections:
        margin = Cm(theme.get("page_margin_cm", 2.0))
        section.top_margin = margin
        section.bottom_margin = margin
        section.left_margin = margin
        section.right_margin = margin

    # Default body font
    style = doc.styles["Normal"]
    style.font.name = theme.get("font_body_single", "Calibri")
    style.font.size = Pt(theme.get("body_size_pt", 11))

    primary_rgb = _hex_to_rgb(theme.get("colour_primary", "000000"))
    text_rgb = _hex_to_rgb(theme.get("colour_text", "000000"))
    muted_rgb = _hex_to_rgb(theme.get("colour_muted", "666666"))

    # Cover page
    if theme.get("show_cover_page", False):
        cover_title = doc.add_paragraph()
        cover_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = cover_title.add_run(title)
        run.font.name = theme.get("font_heading_single", "Calibri")
        run.font.size = Pt(32)
        run.font.bold = True
        run.font.color.rgb = RGBColor(*primary_rgb)

        if subtitle:
            sub = doc.add_paragraph()
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            srun = sub.add_run(subtitle)
            srun.font.size = Pt(14)
            srun.font.color.rgb = RGBColor(*muted_rgb)

        doc.add_paragraph()
        meta = doc.add_paragraph()
        meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_text = datetime.now().strftime("%d %B %Y")
        if author:
            meta_text = f"{author}  |  {meta_text}"
        mrun = meta.add_run(meta_text)
        mrun.font.size = Pt(11)
        mrun.font.color.rgb = RGBColor(*muted_rgb)

        doc.add_page_break()
    else:
        # Inline title for minimal theme
        title_p = doc.add_paragraph()
        trun = title_p.add_run(title)
        trun.font.name = theme.get("font_heading_single", "Calibri")
        trun.font.size = Pt(theme.get("heading_sizes_pt", [22])[0])
        trun.font.bold = True
        trun.font.color.rgb = RGBColor(*text_rgb)
        if subtitle:
            sub = doc.add_paragraph()
            srun = sub.add_run(subtitle)
            srun.font.size = Pt(theme.get("body_size_pt", 11))
            srun.font.color.rgb = RGBColor(*muted_rgb)

    # Body content
    blocks_rendered = 0
    for block in content:
        btype = block.get("type", "paragraph")
        try:
            if btype == "heading":
                _add_heading(doc, block, theme, primary_rgb, text_rgb)
            elif btype == "paragraph":
                _add_paragraph(doc, block, theme, text_rgb)
            elif btype == "list":
                _add_list(doc, block)
            elif btype == "table":
                _add_table(doc, block, theme, primary_rgb)
            elif btype == "rule":
                _add_rule(doc)
            elif btype == "spacer":
                doc.add_paragraph()
            elif btype == "code":
                _add_code(doc, block, theme)
            else:
                _add_paragraph(doc, {"text": str(block)}, theme, text_rgb)
            blocks_rendered += 1
        except Exception:
            # Skip a malformed block rather than aborting the whole document
            continue

    # Footer with generation date and page numbers
    if theme.get("show_generated_footer", True) or theme.get("show_page_numbers", True):
        for section in doc.sections:
            footer = section.footer
            f_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            f_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if theme.get("show_generated_footer", True):
                gen_run = f_para.add_run(
                    f"Generated {datetime.now().strftime('%d %B %Y')}"
                )
                gen_run.font.size = Pt(theme.get("small_size_pt", 9))
                gen_run.font.color.rgb = RGBColor(*muted_rgb)
            if theme.get("show_page_numbers", True):
                if theme.get("show_generated_footer", True):
                    f_para.add_run("   |   ").font.size = Pt(9)
                _add_page_number(f_para, theme.get("small_size_pt", 9), muted_rgb)

    out = Path(output_path)
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    size = out.stat().st_size

    return {
        "path": str(out),
        "size_bytes": size,
        "blocks_rendered": blocks_rendered,
        "theme": theme.get("name", "unknown"),
    }


def _hex_to_rgb(hex_str: str) -> tuple:
    h = (hex_str or "000000").lstrip("#")
    return tuple(int(h[i:i+2], 16) for i in (0, 2, 4))


def _add_heading(doc, block, theme, primary_rgb, text_rgb):
    from docx.shared import Pt, RGBColor
    level = max(1, min(4, int(block.get("level", 1))))
    p = doc.add_paragraph()
    run = p.add_run(block.get("text", ""))
    run.font.name = theme.get("font_heading_single", "Calibri")
    run.font.size = Pt(theme.get("heading_sizes_pt", [22, 16, 13, 11])[level - 1])
    run.font.bold = True
    run.font.color.rgb = RGBColor(*(primary_rgb if level == 1 else text_rgb))


def _add_paragraph(doc, block, theme, text_rgb):
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    run = p.add_run(block.get("text", ""))
    run.font.name = theme.get("font_body_single", "Calibri")
    run.font.size = Pt(theme.get("body_size_pt", 11))
    run.font.color.rgb = RGBColor(*text_rgb)


def _add_list(doc, block):
    items = block.get("items", [])
    ordered = bool(block.get("ordered", False))
    style_name = "List Number" if ordered else "List Bullet"
    for item in items:
        doc.add_paragraph(str(item), style=style_name)


def _add_table(doc, block, theme, primary_rgb):
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.table import WD_TABLE_ALIGNMENT

    headers = block.get("headers", [])
    rows = block.get("rows", [])
    if not headers and not rows:
        return

    n_cols = max(len(headers), max((len(r) for r in rows), default=0))
    n_rows = (1 if headers else 0) + len(rows)
    if n_cols == 0 or n_rows == 0:
        return

    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Light Grid"

    # column_widths: same schema as pdf/xlsx. Values <= 1.0 are relative
    # weights that sum to the printable page width; values > 1.0 are
    # absolute centimetres.
    col_widths_cfg = block.get("column_widths") or []
    if col_widths_cfg and len(col_widths_cfg) == n_cols:
        try:
            values = [float(v) for v in col_widths_cfg]
            total = sum(values)
            if total > 0:
                # Printable width = page width - margins. A4 portrait
                # with 2cm margins = 17cm. Close enough across themes.
                printable_cm = 21.0 - 2 * theme.get("page_margin_cm", 2.0)
                if all(v <= 1.0 for v in values):
                    widths_cm = [v / total * printable_cm for v in values]
                else:
                    widths_cm = values
                for ci, w in enumerate(widths_cm):
                    for row in table.rows:
                        if ci < len(row.cells):
                            row.cells[ci].width = Cm(max(1.0, w))
        except (TypeError, ValueError):
            pass

    # Header row
    if headers:
        hdr_cells = table.rows[0].cells
        header_fill = theme.get("table_header_fill", "0B5FFF")
        header_text = _hex_to_rgb(theme.get("table_header_text", "FFFFFF"))
        for i, h in enumerate(headers[:n_cols]):
            cell = hdr_cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(h))
            run.bold = True
            run.font.color.rgb = RGBColor(*header_text)
            run.font.size = Pt(theme.get("body_size_pt", 11))
            _set_cell_fill(cell, header_fill)

    # Body rows
    body_offset = 1 if headers else 0
    zebra = theme.get("table_zebra", False)
    zebra_fill = theme.get("colour_zebra", "F8FAFC")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row[:n_cols]):
            cell = table.rows[ri + body_offset].cells[ci]
            cell.text = str(val)
            if zebra and ri % 2 == 1:
                _set_cell_fill(cell, zebra_fill)


def _set_cell_fill(cell, hex_colour: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour.lstrip("#"))
    tc_pr.append(shd)


def _add_rule(doc):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_code(doc, block, theme):
    from docx.shared import Pt, RGBColor
    p = doc.add_paragraph()
    run = p.add_run(block.get("text", ""))
    run.font.name = theme.get("font_mono_single", "Consolas")
    run.font.size = Pt(theme.get("small_size_pt", 9))
    run.font.color.rgb = RGBColor(*_hex_to_rgb(theme.get("colour_text", "000000")))
    _set_paragraph_fill(p, theme.get("colour_zebra", "F8FAFC"))


def _set_paragraph_fill(paragraph, hex_colour: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_colour.lstrip("#"))
    p_pr.append(shd)


def _add_page_number(paragraph, size_pt: int, rgb: tuple):
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    run = paragraph.add_run("Page ")
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(*rgb)

    fld_run = paragraph.add_run()
    fld_run.font.size = Pt(size_pt)
    fld_run.font.color.rgb = RGBColor(*rgb)
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    fld_run._r.append(fld_char_begin)
    fld_run._r.append(instr)
    fld_run._r.append(fld_char_end)
