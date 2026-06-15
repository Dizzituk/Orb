# FILE: app/documents/docx_convert.py
# Purpose: docx ⇄ Univer doc snapshot (python-docx): paragraphs, headings,
#          bold/italic, bullet/numbered lists. Lossy by design.
# Called-by: app.documents.router
# Depends-on: python-docx
# Last-renovated: 2026-06-12
"""
docx conversion.

WHAT SURVIVES the round trip (v1, by design):
  - paragraph text and order
  - headings 1-3 (encoded as bold + larger font in Univer; recovered on
    save by the size heuristic below)
  - bold / italic runs
  - bullet and numbered lists ("• " / "1. " text prefixes in the editor;
    mapped back to Word's List Bullet / List Number styles on save)
WHAT IS DROPPED (documented loss — complex Word features round-trip
best-effort): tables, images, footnotes, comments, tracked changes,
hyperlink targets (text kept), nested list levels, fonts/colours, page
layout. A complex document still OPENS — unknown structures degrade to
plain paragraphs and a warning is logged, never a crash.

Univer doc snapshot primer: body.dataStream is the whole text with '\\r'
between paragraphs and a final '\\r\\n'; textRuns are [st, ed) character
ranges with ts.bl/ts.it/ts.fs; paragraphs[].startIndex points AT each
'\\r'. Heading sizes: H1 22pt, H2 18pt, H3 15pt (body 11pt) — the save
path inverts exactly these.
"""
from __future__ import annotations

import logging
from pathlib import Path

logger = logging.getLogger(__name__)

H1_SIZE, H2_SIZE, H3_SIZE, BODY_SIZE = 22, 18, 15, 11
BULLET_PREFIX = "• "

# Univer BooleanNumber.TRUE
_TRUE = 1


def _heading_level(style_name: str) -> int:
    name = (style_name or "").lower()
    if name.startswith("heading"):
        try:
            return min(3, max(1, int(name.split()[-1])))
        except Exception:
            return 1
    if name == "title":
        return 1
    return 0


def _list_kind(style_name: str) -> str:
    name = (style_name or "").lower()
    if "list bullet" in name:
        return "bullet"
    if "list number" in name:
        return "number"
    return ""


def docx_to_snapshot(path: str) -> dict:
    """python-docx read -> Univer IDocumentData-shaped dict (never raises
    on weird content — degrades to plain text + warning)."""
    import docx

    document = docx.Document(path)
    stream_parts: list[str] = []
    text_runs: list[dict] = []
    paragraphs: list[dict] = []
    cursor = 0
    number_counter = 0

    for para in document.paragraphs:
        try:
            style_name = para.style.name if para.style is not None else ""
        except Exception:
            style_name = ""
        heading = _heading_level(style_name)
        list_kind = _list_kind(style_name)

        prefix = ""
        if list_kind == "bullet":
            prefix = BULLET_PREFIX
            number_counter = 0
        elif list_kind == "number":
            number_counter += 1
            prefix = f"{number_counter}. "
        else:
            number_counter = 0

        if prefix:
            stream_parts.append(prefix)
            cursor += len(prefix)

        para_has_text = False
        for run in para.runs:
            text = run.text or ""
            if not text:
                continue
            para_has_text = True
            start = cursor
            stream_parts.append(text)
            cursor += len(text)
            ts: dict = {}
            if heading:
                ts["bl"] = _TRUE
                ts["fs"] = {1: H1_SIZE, 2: H2_SIZE, 3: H3_SIZE}[heading]
            else:
                if run.bold:
                    ts["bl"] = _TRUE
                if run.italic:
                    ts["it"] = _TRUE
            if ts:
                text_runs.append({"st": start, "ed": cursor, "ts": ts})
        if not para_has_text and para.text:
            # runs missing (unusual XML) — fall back to the plain text
            stream_parts.append(para.text)
            cursor += len(para.text)

        stream_parts.append("\r")
        paragraphs.append({"startIndex": cursor})
        cursor += 1

    stream_parts.append("\n")
    data_stream = "".join(stream_parts)

    skipped = len(document.tables)
    if skipped:
        logger.warning("[documents] %s: %d table(s) not shown in the editor "
                       "(v1 loss policy)", path, skipped)

    return {
        "id": f"doc-{Path(path).stem[:40]}",
        "body": {
            "dataStream": data_stream,
            "textRuns": text_runs,
            "paragraphs": paragraphs,
            "sectionBreaks": [{"startIndex": max(len(data_stream) - 1, 0)}],
        },
        "documentStyle": {
            "pageSize": {"width": 595, "height": 842},
            "marginTop": 50, "marginBottom": 50,
            "marginLeft": 50, "marginRight": 50,
        },
    }


def _runs_for_span(text_runs: list, start: int, end: int) -> list:
    """textRuns clipped to [start, end) of one paragraph."""
    out = []
    for run in text_runs or []:
        st, ed = int(run.get("st", 0)), int(run.get("ed", 0))
        if ed <= start or st >= end:
            continue
        out.append({"st": max(st, start), "ed": min(ed, end),
                    "ts": run.get("ts") or {}})
    return out


def snapshot_to_docx(snapshot: dict, path: str) -> None:
    """Univer doc snapshot -> .docx (inverts the read mapping above)."""
    import docx

    body = snapshot.get("body") or {}
    stream: str = body.get("dataStream") or ""
    text_runs = body.get("textRuns") or []
    document = docx.Document()

    cursor = 0
    for para_meta in body.get("paragraphs") or []:
        end = int(para_meta.get("startIndex", cursor))
        text = stream[cursor:end]
        runs = _runs_for_span(text_runs, cursor, end)

        # invert the heading encoding: whole-paragraph bold at known sizes
        sizes = [int((r.get("ts") or {}).get("fs") or 0) for r in runs]
        all_bold = bool(runs) and all((r.get("ts") or {}).get("bl") == _TRUE for r in runs)
        covers_all = bool(runs) and runs[0]["st"] <= cursor and runs[-1]["ed"] >= end
        heading = 0
        if text.strip() and all_bold and covers_all and sizes:
            biggest = max(sizes)
            if biggest >= H1_SIZE:
                heading = 1
            elif biggest >= H2_SIZE:
                heading = 2
            elif biggest >= H3_SIZE:
                heading = 3

        style = None
        body_text = text
        if heading:
            paragraph = document.add_heading(body_text.strip(), level=heading)
            cursor = end + 1
            continue
        if text.startswith(BULLET_PREFIX):
            style = "List Bullet"
            body_text = text[len(BULLET_PREFIX):]
        else:
            import re
            numbered = re.match(r"^(\d+)\.\s", text)
            if numbered:
                style = "List Number"
                body_text = text[numbered.end():]

        paragraph = document.add_paragraph(style=style)
        offset = end - len(body_text)  # stream index where body_text begins
        pos = offset
        for run_meta in _runs_for_span(text_runs, offset, end):
            st, ed = run_meta["st"], run_meta["ed"]
            if st > pos:
                paragraph.add_run(stream[pos:st])
            piece = paragraph.add_run(stream[st:ed])
            ts = run_meta.get("ts") or {}
            if ts.get("bl") == _TRUE:
                piece.bold = True
            if ts.get("it") == _TRUE:
                piece.italic = True
            pos = ed
        if pos < end:
            paragraph.add_run(stream[pos:end])

        cursor = end + 1

    document.save(path)
