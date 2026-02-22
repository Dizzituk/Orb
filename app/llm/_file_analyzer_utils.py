import io
import json
import logging
import re
from typing import Any, Callable, Dict, Optional, Tuple
logger = logging.getLogger(__name__)
logger = logging.getLogger(__name__)


def _analyze_with_pymupdf(
    pdf_path: Optional[str],
    pdf_bytes: Optional[bytes],
    result: Dict[str, Any]
) -> Dict[str, Any]:
    """Analyze PDF using PyMuPDF (fitz) - most accurate image detection."""
    import fitz
    
    if pdf_bytes:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    else:
        doc = fitz.open(pdf_path)
    
    try:
        result["page_count"] = len(doc)
        total_text = ""
        image_count = 0
        
        for page in doc:
            # Extract text
            total_text += page.get_text()
            
            # Count images on this page
            image_list = page.get_images(full=True)
            image_count += len(image_list)
        
        result["text_chars"] = len(total_text)
        result["image_count"] = image_count
        
        if result["page_count"] > 0:
            result["chars_per_page"] = result["text_chars"] // result["page_count"]
        
        return result
        
    finally:
        doc.close()

def _extract_docx_text(
    file_path: Optional[str],
    file_bytes: Optional[bytes]
) -> Tuple[str, Optional[str]]:
    """Extract text from DOCX."""
    try:
        from docx import Document
        
        if file_bytes:
            doc = Document(io.BytesIO(file_bytes))
        else:
            doc = Document(file_path)
        
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text.strip(), None
    except ImportError:
        return "", "python-docx not installed"
    except Exception as e:
        return "", f"DOCX extraction failed: {str(e)}"

def prepare_attachment_info(file_path: str) -> Dict[str, Any]:
    """
    Prepare attachment info dict for job_classifier.
    This returns a dict compatible with AttachmentInfo schema.
    """
    info = get_file_info(file_path)
    
    return {
        "filename": info["filename"],
        "size_bytes": info["size_bytes"],
        "mime_type": _guess_mime_type(info["extension"]),
        "pdf_image_count": info.get("pdf_image_count"),
        "pdf_text_chars": info.get("pdf_text_chars"),
        "pdf_page_count": info.get("pdf_page_count"),
    }

def _guess_mime_type(extension: str) -> Optional[str]:
    """Guess MIME type from extension."""
    mime_map = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".gif": "image/gif",
        ".webp": "image/webp",
        ".mp4": "video/mp4",
        ".mov": "video/quicktime",
        ".avi": "video/x-msvideo",
        ".py": "text/x-python",
        ".js": "application/javascript",
        ".ts": "application/typescript",
        ".json": "application/json",
        ".yaml": "application/x-yaml",
        ".yml": "application/x-yaml",
    }
    return mime_map.get(extension.lower())

def extract_text_content(
    file_path: str,
    mime_type: Optional[str] = None,
) -> Optional[str]:
    """
    Extract text content from a file.
    
    Called from main.py as: extract_text_content(str(file_path), mime_type)
    
    Args:
        file_path: Path to the file
        mime_type: MIME type (used for routing, not extraction)
    
    Returns:
        Extracted text string, or None if extraction failed
    """
    text, error = extract_text(file_path=file_path)
    if error:
        logger.warning(f"Text extraction error for {file_path}: {error}")
    return text if text else None

def is_pdf_mime_type(mime_type: Optional[str]) -> bool:
    """Check if MIME type is PDF."""
    if not mime_type:
        return False
    return mime_type == "application/pdf"

def parse_cv_with_llm(
    raw_text: Optional[str],
    filename: str,
    llm_call: Callable[[str], str],
) -> dict:
    """
    Parse a CV/resume using LLM.
    
    Called from main.py as: parse_cv_with_llm(raw_text, original_name, simple_llm_call)
    
    Args:
        raw_text: Already-extracted text content from the CV (may be None)
        filename: Original filename for context
        llm_call: Callable that takes a prompt string and returns response string
    
    Returns:
        dict with parsed CV fields (name, email, skills, experience, etc.)
    """
    # Handle None, empty, or non-string input
    if not raw_text or not isinstance(raw_text, str) or not raw_text.strip():
        return {
            "name": None,
            "email": None,
            "phone": None,
            "skills": [],
            "roles": [],
            "education": [],
            "error": "No text content to parse",
        }
    
    # Use LLM to extract structured data
    prompt = f"""Extract structured information from this CV/resume. Return JSON with these fields:
- name: Full name
- email: Email address (if found)
- phone: Phone number (if found)
- skills: List of skills mentioned
- roles: List of work experiences (each with title, company, dates if available)
- education: List of education entries

CV Content:
{raw_text[:8000]}

Respond with only valid JSON, no other text."""

    try:
        response = llm_call(prompt)
        # Try to parse JSON from response
        import json
        # Find JSON in response (might be wrapped in markdown)
        json_match = re.search(r'\{[\s\S]*\}', response)
        if json_match:
            return json.loads(json_match.group())
    except Exception as e:
        logger.warning(f"CV parsing failed: {e}")
    
    # Fallback: basic extraction
    return {
        "name": None,
        "email": None,
        "phone": None,
        "skills": [],
        "roles": [],
        "education": [],
        "raw_text": raw_text[:1000],
        "error": "LLM parsing failed, raw text preserved",
    }

def generate_document_summary(
    raw_text: Optional[str],
    filename: str,
    doc_type: str,
    llm_call: Callable[[str], str],
    max_length: int = 500,
) -> str:
    """
    Generate a summary of a document using LLM.
    
    Called from main.py as: generate_document_summary(raw_text, original_name, doc_type, simple_llm_call)
    
    Args:
        raw_text: Already-extracted text content (may be None)
        filename: Original filename for context
        doc_type: Type of document (pdf, docx, text, etc.)
        llm_call: Callable that takes a prompt string and returns response string
        max_length: Target max length for summary
    
    Returns:
        Summary string
    """
    # Handle None, empty string, or whitespace-only
    if not raw_text or not isinstance(raw_text, str) or not raw_text.strip():
        return f"Document uploaded: {filename}"
    
    # Truncate very long documents for the prompt
    text_for_prompt = raw_text[:6000] if len(raw_text) > 6000 else raw_text
    
    prompt = f"""Summarize this {doc_type} document in 2-3 sentences. Be concise and focus on the main topics/purpose.

Document: {filename}

Content:
{text_for_prompt}

Summary:"""

    try:
        response = llm_call(prompt)
        if response and response.strip():
            # Clean up response
            summary = response.strip()
            # Remove any "Summary:" prefix if LLM included it
            if summary.lower().startswith("summary:"):
                summary = summary[8:].strip()
            return summary[:max_length] if len(summary) > max_length else summary
    except Exception as e:
        logger.warning(f"Document summary generation failed: {e}")
    
    # Fallback: return truncated text
    if len(raw_text) > max_length:
        return raw_text[:max_length].strip() + "..."
    return raw_text.strip()
