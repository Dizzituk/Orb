# FILE: app/llm/file_analyzer.py
"""
File analysis utilities for LLM routing.

Version: 0.13.0 - Phase 4 Routing Fix

Key functions:
- analyze_pdf_content(): Count images and extract text for routing decisions
- extract_text(): Extract text from various file types
- is_binary_file(): Detect binary vs text files

PDF ROUTING RULE:
- image_count == 0 → GPT text.heavy
- image_count > 0 → Gemini image.complex
"""

import os
import io
import re
import json
import logging
from typing import Optional, Dict, Any, Tuple, Callable
from pathlib import Path

logger = logging.getLogger(__name__)


def analyze_pdf_content(
    pdf_path: Optional[str] = None,
    pdf_bytes: Optional[bytes] = None
) -> Dict[str, Any]:
    """
    Analyze PDF content for routing decisions.
    
    Counts embedded images and extracts text to determine
    whether PDF should route to GPT (text-only) or Gemini (has images).
    
    Args:
        pdf_path: Path to PDF file
        pdf_bytes: Raw PDF bytes (alternative to path)
    
    Returns:
        dict with keys:
        - image_count: Number of embedded images (0 = text-only)
        - text_chars: Total character count from extracted text
        - page_count: Number of pages
        - chars_per_page: Average chars per page
        - error: Error message if analysis failed
    
    Routing rule:
    - image_count == 0 → route to GPT text.heavy
    - image_count > 0 → route to Gemini image.complex
    """
    result = {
        "image_count": 0,
        "text_chars": 0,
        "page_count": 0,
        "chars_per_page": 0,
        "error": None,
    }
    
    try:
        # Try PyMuPDF (fitz) first - best for image detection
        try:
            import fitz  # PyMuPDF
            return _analyze_with_pymupdf(pdf_path, pdf_bytes, result)
        except ImportError:
            pass
        
        # Fall back to pypdf
        try:
            from pypdf import PdfReader
            return _analyze_with_pypdf(pdf_path, pdf_bytes, result)
        except ImportError:
            pass
        
        # Fall back to PyPDF2
        try:
            import PyPDF2
            return _analyze_with_pypdf2(pdf_path, pdf_bytes, result)
        except ImportError:
            pass
        
        result["error"] = "No PDF library available (install pymupdf, pypdf, or PyPDF2)"
        logger.error(result["error"])
        return result
        
    except Exception as e:
        result["error"] = f"PDF analysis failed: {str(e)}"
        logger.error(result["error"], exc_info=True)
        return result


def _analyze_with_pymupdf(
    pdf_path: Optional[str],
    pdf_bytes: Optional[bytes],
    result: Dict[str, Any]
) -> Dict[str, Any]:
    """Analyze PDF using PyMuPDF (fitz) - most accurate image detection."""
    import fitz
    
    if pdf_bytes:
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
    else:
        doc = fitz.open(pdf_path)
    
    try:
        result["page_count"] = len(doc)
        total_text = ""
        image_count = 0
        
        for page in doc:
            # Extract text
            total_text += page.get_text()
            
            # Count images on this page
            image_list = page.get_images(full=True)
            image_count += len(image_list)
        
        result["text_chars"] = len(total_text)
        result["image_count"] = image_count
        
        if result["page_count"] > 0:
            result["chars_per_page"] = result["text_chars"] // result["page_count"]
        
        return result
        
    finally:
        doc.close()


def _analyze_with_pypdf(
    pdf_path: Optional[str],
    pdf_bytes: Optional[bytes],
    result: Dict[str, Any]
) -> Dict[str, Any]:
    """Analyze PDF using pypdf."""
    from pypdf import PdfReader
    
    if pdf_bytes:
        reader = PdfReader(io.BytesIO(pdf_bytes))
    else:
        reader = PdfReader(pdf_path)
    
    result["page_count"] = len(reader.pages)
    total_text = ""
    image_count = 0
    
    for page in reader.pages:
        # Extract text
        text = page.extract_text() or ""
        total_text += text
        
        # Count images (pypdf stores them in /Resources/XObject)
        try:
            if "/Resources" in page:
                resources = page["/Resources"]
                if "/XObject" in resources:
                    xobjects = resources["/XObject"]
                    if hasattr(xobjects, "get_object"):
                        xobjects = xobjects.get_object()
                    for obj_name in xobjects:
                        obj = xobjects[obj_name]
                        if hasattr(obj, "get_object"):
                            obj = obj.get_object()
                        if obj.get("/Subtype") == "/Image":
                            image_count += 1
        except Exception as e:
            logger.debug(f"Error counting images on page: {e}")
    
    result["text_chars"] = len(total_text)
    result["image_count"] = image_count
    
    if result["page_count"] > 0:
        result["chars_per_page"] = result["text_chars"] // result["page_count"]
    
    return result


def _analyze_with_pypdf2(
    pdf_path: Optional[str],
    pdf_bytes: Optional[bytes],
    result: Dict[str, Any]
) -> Dict[str, Any]:
    """Analyze PDF using PyPDF2 (legacy fallback)."""
    import PyPDF2
    
    if pdf_bytes:
        reader = PyPDF2.PdfReader(io.BytesIO(pdf_bytes))
    else:
        reader = PyPDF2.PdfReader(pdf_path)
    
    result["page_count"] = len(reader.pages)
    total_text = ""
    image_count = 0
    
    for page in reader.pages:
        # Extract text
        text = page.extract_text() or ""
        total_text += text
        
        # Count images
        try:
            if "/Resources" in page:
                resources = page["/Resources"].get_object()
                if "/XObject" in resources:
                    xobjects = resources["/XObject"].get_object()
                    for obj_name in xobjects:
                        obj = xobjects[obj_name].get_object()
                        if obj.get("/Subtype") == "/Image":
                            image_count += 1
        except Exception as e:
            logger.debug(f"Error counting images on page: {e}")
    
    result["text_chars"] = len(total_text)
    result["image_count"] = image_count
    
    if result["page_count"] > 0:
        result["chars_per_page"] = result["text_chars"] // result["page_count"]
    
    return result


def extract_text(
    file_path: Optional[str] = None,
    file_bytes: Optional[bytes] = None,
    filename: Optional[str] = None
) -> Tuple[str, Optional[str]]:
    """
    Extract text from a file.
    
    Args:
        file_path: Path to file
        file_bytes: Raw file bytes
        filename: Original filename (for extension detection)
    
    Returns:
        (text, error) - extracted text and optional error message
    """
    if not file_path and not file_bytes:
        return "", "No file provided"
    
    # Determine extension
    if file_path:
        ext = Path(file_path).suffix.lower()
    elif filename:
        ext = Path(filename).suffix.lower()
    else:
        return "", "Cannot determine file type"
    
    try:
        if ext == ".pdf":
            return _extract_pdf_text(file_path, file_bytes)
        
        elif ext == ".docx":
            return _extract_docx_text(file_path, file_bytes)
        
        elif ext in {".txt", ".md", ".py", ".js", ".ts", ".json", ".yaml", ".yml"}:
            if file_bytes:
                return file_bytes.decode("utf-8", errors="replace"), None
            else:
                with open(file_path, "r", encoding="utf-8", errors="replace") as f:
                    return f.read(), None
        
        else:
            return "", f"Unsupported file type: {ext}"
            
    except Exception as e:
        return "", f"Text extraction failed: {str(e)}"


def _extract_pdf_text(
    file_path: Optional[str],
    file_bytes: Optional[bytes]
) -> Tuple[str, Optional[str]]:
    """Extract text from PDF."""
    try:
        import fitz
        if file_bytes:
            doc = fitz.open(stream=file_bytes, filetype="pdf")
        else:
            doc = fitz.open(file_path)
        
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text, None
    except ImportError:
        pass
    
    try:
        from pypdf import PdfReader
        if file_bytes:
            reader = PdfReader(io.BytesIO(file_bytes))
        else:
            reader = PdfReader(file_path)
        
        text = ""
        for page in reader.pages:
            text += page.extract_text() or ""
        return text, None
    except ImportError:
        pass
    
    return "", "No PDF library available"


def _extract_docx_text(
    file_path: Optional[str],
    file_bytes: Optional[bytes]
) -> Tuple[str, Optional[str]]:
    """Extract text from DOCX."""
    try:
        from docx import Document
        
        if file_bytes:
            doc = Document(io.BytesIO(file_bytes))
        else:
            doc = Document(file_path)
        
        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"
        return text.strip(), None
    except ImportError:
        return "", "python-docx not installed"
    except Exception as e:
        return "", f"DOCX extraction failed: {str(e)}"


def is_binary_file(
    file_path: Optional[str] = None,
    file_bytes: Optional[bytes] = None,
    chunk_size: int = 8192
) -> bool:
    """
    Detect if a file is binary (not text).
    Uses null byte detection - presence of null bytes indicates binary.
    """
    if file_bytes:
        chunk = file_bytes[:chunk_size]
    elif file_path:
        try:
            with open(file_path, "rb") as f:
                chunk = f.read(chunk_size)
        except Exception:
            return True
    else:
        return False
    
    return b"\x00" in chunk


def get_file_info(file_path: str) -> Dict[str, Any]:
    """
    Get file information for routing decisions.
    
    Args:
        file_path: Path to file
    
    Returns:
        dict with filename, size_bytes, extension, is_binary,
        and for PDFs: pdf_image_count, pdf_text_chars, pdf_page_count
    """
    path = Path(file_path)
    
    info = {
        "filename": path.name,
        "size_bytes": path.stat().st_size if path.exists() else 0,
        "extension": path.suffix.lower(),
        "is_binary": is_binary_file(file_path=file_path),
    }
    
    # For PDFs, add analysis
    if info["extension"] == ".pdf":
        pdf_analysis = analyze_pdf_content(pdf_path=file_path)
        info["pdf_image_count"] = pdf_analysis["image_count"]
        info["pdf_text_chars"] = pdf_analysis["text_chars"]
        info["pdf_page_count"] = pdf_analysis["page_count"]
        if pdf_analysis["error"]:
            info["pdf_error"] = pdf_analysis["error"]
    
    return info


def prepare_attachment_info(file_path: str) -> Dict[str, Any]:
    """
    Prepare attachment info dict for job_classifier.
    This returns a dict compatible with AttachmentInfo schema.
    """
    info = get_file_info(file_path)
    
    return {
        "filename": info["filename"],
        "size_bytes": info["size_bytes"],
        "mime_type": _guess_mime_type(info["extension"]),
        "pdf_image_count": info.get("pdf_image_count"),
        "pdf_text_chars": info.get("pdf_text_chars"),
        "pdf_page_count": info.get("pdf_page_count"),
    }


def _guess_mime_type(extension: str) -> Optional[str]:
    """Guess MIME type from extension."""
    mime_map = {
        ".pdf": "application/pdf",
        ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".doc": "application/msword",
        ".txt": "text/plain",
        ".md": "text/markdown",
        ".png": "image/png",
        ".jpg": "image/jpeg",
        ".jpeg": "image/jpeg",
        ".gif": "image/gif",
        ".webp": "image/webp",
        ".mp4": "video/mp4",
        ".mov": "video/quicktime",
        ".avi": "video/x-msvideo",
        ".py": "text/x-python",
        ".js": "application/javascript",
        ".ts": "application/typescript",
        ".json": "application/json",
        ".yaml": "application/x-yaml",
        ".yml": "application/x-yaml",
    }
    return mime_map.get(extension.lower())


# =============================================================================
# BACKWARD COMPATIBILITY ALIASES
# =============================================================================

def extract_text_content(
    file_path: str,
    mime_type: Optional[str] = None,
) -> Optional[str]:
    """
    Extract text content from a file.
    
    Called from main.py as: extract_text_content(str(file_path), mime_type)
    
    Args:
        file_path: Path to the file
        mime_type: MIME type (used for routing, not extraction)
    
    Returns:
        Extracted text string, or None if extraction failed
    """
    text, error = extract_text(file_path=file_path)
    if error:
        logger.warning(f"Text extraction error for {file_path}: {error}")
    return text if text else None


# =============================================================================
# MIME TYPE DETECTION
# =============================================================================

def is_video_mime_type(mime_type: Optional[str]) -> bool:
    """Check if MIME type is a video format."""
    if not mime_type:
        return False
    return mime_type.startswith("video/") or mime_type in {
        "application/x-matroska",
        "application/vnd.rn-realmedia",
    }


def is_audio_mime_type(mime_type: Optional[str]) -> bool:
    """Check if MIME type is an audio format."""
    if not mime_type:
        return False
    return mime_type.startswith("audio/") or mime_type in {
        "application/ogg",
    }


def is_image_mime_type(mime_type: Optional[str]) -> bool:
    """Check if MIME type is an image format."""
    if not mime_type:
        return False
    return mime_type.startswith("image/")


def is_pdf_mime_type(mime_type: Optional[str]) -> bool:
    """Check if MIME type is PDF."""
    if not mime_type:
        return False
    return mime_type == "application/pdf"


# =============================================================================
# DOCUMENT TYPE DETECTION
# =============================================================================

def detect_document_type(
    content_or_path: Optional[str] = None,
    filename: Optional[str] = None,
    mime_type: Optional[str] = None,
) -> str:
    """
    Detect document type from content, filename, or MIME type.
    
    Called from main.py as: detect_document_type(raw_text, original_name)
    
    Args:
        content_or_path: Either file content (text) or file path - used for fallback heuristics
        filename: Original filename with extension
        mime_type: Optional MIME type hint
    
    Returns one of:
    - "pdf"
    - "docx"
    - "image"
    - "video"
    - "audio"
    - "code"
    - "text"
    - "unknown"
    """
    # Determine extension from filename
    ext = None
    if filename:
        try:
            ext = Path(filename).suffix.lower()
        except (TypeError, ValueError):
            # filename might not be a valid path string
            if isinstance(filename, str) and "." in filename:
                ext = "." + filename.rsplit(".", 1)[-1].lower()
    
    # Check MIME type first
    if mime_type:
        if is_pdf_mime_type(mime_type):
            return "pdf"
        if mime_type in {
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "application/msword",
        }:
            return "docx"
        if is_image_mime_type(mime_type):
            return "image"
        if is_video_mime_type(mime_type):
            return "video"
        if is_audio_mime_type(mime_type):
            return "audio"
    
    # Fall back to extension
    if ext:
        if ext == ".pdf":
            return "pdf"
        if ext in {".docx", ".doc"}:
            return "docx"
        if ext in {".png", ".jpg", ".jpeg", ".gif", ".webp", ".bmp", ".tiff"}:
            return "image"
        if ext in {".mp4", ".mov", ".avi", ".mkv", ".webm", ".m4v"}:
            return "video"
        if ext in {".mp3", ".wav", ".ogg", ".flac", ".m4a", ".aac"}:
            return "audio"
        if ext in {".py", ".js", ".ts", ".java", ".cpp", ".c", ".h", ".go", ".rs", ".rb"}:
            return "code"
        if ext in {".txt", ".md", ".rst", ".csv", ".json", ".yaml", ".yml", ".xml", ".html"}:
            return "text"
    
    # If we have content, try to detect from content
    if content_or_path and isinstance(content_or_path, str):
        # Check if it looks like a file path
        if len(content_or_path) < 500 and ("/" in content_or_path or "\\" in content_or_path):
            try:
                path_ext = Path(content_or_path).suffix.lower()
                if path_ext:
                    # Recurse with the path as filename
                    return detect_document_type(filename=content_or_path)
            except (TypeError, ValueError):
                pass
    
    return "unknown"


# =============================================================================
# LLM-BASED DOCUMENT PROCESSING (stubs for backward compat)
# =============================================================================

def parse_cv_with_llm(
    raw_text: Optional[str],
    filename: str,
    llm_call: Callable[[str], str],
) -> dict:
    """
    Parse a CV/resume using LLM.
    
    Called from main.py as: parse_cv_with_llm(raw_text, original_name, simple_llm_call)
    
    Args:
        raw_text: Already-extracted text content from the CV (may be None)
        filename: Original filename for context
        llm_call: Callable that takes a prompt string and returns response string
    
    Returns:
        dict with parsed CV fields (name, email, skills, experience, etc.)
    """
    # Handle None, empty, or non-string input
    if not raw_text or not isinstance(raw_text, str) or not raw_text.strip():
        return {
            "name": None,
            "email": None,
            "phone": None,
            "skills": [],
            "roles": [],
            "education": [],
            "error": "No text content to parse",
        }
    
    # Use LLM to extract structured data
    prompt = f"""Extract structured information from this CV/resume. Return JSON with these fields:
- name: Full name
- email: Email address (if found)
- phone: Phone number (if found)
- skills: List of skills mentioned
- roles: List of work experiences (each with title, company, dates if available)
- education: List of education entries

CV Content:
{raw_text[:8000]}

Respond with only valid JSON, no other text."""

    try:
        response = llm_call(prompt)
        # Try to parse JSON from response
        import json
        # Find JSON in response (might be wrapped in markdown)
        json_match = re.search(r'\{[\s\S]*\}', response)
        if json_match:
            return json.loads(json_match.group())
    except Exception as e:
        logger.warning(f"CV parsing failed: {e}")
    
    # Fallback: basic extraction
    return {
        "name": None,
        "email": None,
        "phone": None,
        "skills": [],
        "roles": [],
        "education": [],
        "raw_text": raw_text[:1000],
        "error": "LLM parsing failed, raw text preserved",
    }


def generate_document_summary(
    raw_text: Optional[str],
    filename: str,
    doc_type: str,
    llm_call: Callable[[str], str],
    max_length: int = 500,
) -> str:
    """
    Generate a summary of a document using LLM.
    
    Called from main.py as: generate_document_summary(raw_text, original_name, doc_type, simple_llm_call)
    
    Args:
        raw_text: Already-extracted text content (may be None)
        filename: Original filename for context
        doc_type: Type of document (pdf, docx, text, etc.)
        llm_call: Callable that takes a prompt string and returns response string
        max_length: Target max length for summary
    
    Returns:
        Summary string
    """
    # Handle None, empty string, or whitespace-only
    if not raw_text or not isinstance(raw_text, str) or not raw_text.strip():
        return f"Document uploaded: {filename}"
    
    # Truncate very long documents for the prompt
    text_for_prompt = raw_text[:6000] if len(raw_text) > 6000 else raw_text
    
    prompt = f"""Summarize this {doc_type} document in 2-3 sentences. Be concise and focus on the main topics/purpose.

Document: {filename}

Content:
{text_for_prompt}

Summary:"""

    try:
        response = llm_call(prompt)
        if response and response.strip():
            # Clean up response
            summary = response.strip()
            # Remove any "Summary:" prefix if LLM included it
            if summary.lower().startswith("summary:"):
                summary = summary[8:].strip()
            return summary[:max_length] if len(summary) > max_length else summary
    except Exception as e:
        logger.warning(f"Document summary generation failed: {e}")
    
    # Fallback: return truncated text
    if len(raw_text) > max_length:
        return raw_text[:max_length].strip() + "..."
    return raw_text.strip()