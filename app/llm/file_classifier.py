# FILE: app/llm/file_classifier.py
"""
File Classification and File Map Builder for Orb Routing.

Version: 1.0.0 - Critical Pipeline Spec Implementation

Implements Spec §1 (Attachment Classification) and §2 (File Map):
- Classifies each attachment into exactly ONE bucket
- Builds stable [FILE_X] naming scheme
- Computes canonical modality flags

FILE TYPE BUCKETS (exactly one per file):
- TEXT_FILE: Pure text only (no embedded images) - .txt, .md, .docx (no images), .pdf (no images)
- CODE_FILE: Programming code - .py, .js, .ts, .java, .cpp, etc.
- IMAGE_FILE: Standalone images - .png, .jpg, .jpeg, .webp, static .gif
- VIDEO_FILE: Video formats - .mp4, .mov, .webm, .avi
- MIXED_FILE: Files containing both text and images - .pdf with images, .docx with images

MODALITY FLAGS:
- HAS_TEXT = len(text_files) > 0
- HAS_CODE = len(code_files) > 0
- HAS_IMAGE = len(image_files) > 0 OR len(mixed_files) > 0
- HAS_VIDEO = len(video_files) > 0
- HAS_MIXED = len(mixed_files) > 0

FILE MAP FORMAT:
[FILE_1] original_filename.ext (FILE_TYPE)
[FILE_2] another_file.pdf (MIXED_FILE: text + images)
...

Usage:
    from app.llm.file_classifier import classify_attachments, build_file_map
    
    result = classify_attachments(attachments)
    file_map_str = build_file_map(result.classified_files)
"""

import os
import io
import logging
from enum import Enum
from pathlib import Path
from typing import Optional, List, Dict, Any, Tuple
from dataclasses import dataclass, field
from pydantic import BaseModel

logger = logging.getLogger(__name__)


# =============================================================================
# FILE TYPE ENUM (Spec §1.1)
# =============================================================================

class FileType(str, Enum):
    """
    Canonical file type buckets.
    
    Each attachment is classified into exactly ONE of these types.
    """
    TEXT_FILE = "TEXT_FILE"      # Pure text only (no embedded images)
    CODE_FILE = "CODE_FILE"      # Programming code
    IMAGE_FILE = "IMAGE_FILE"    # Standalone images
    VIDEO_FILE = "VIDEO_FILE"    # Video formats
    MIXED_FILE = "MIXED_FILE"    # Text + images (e.g., PDF with images, docx with diagrams)


# =============================================================================
# FILE EXTENSIONS
# =============================================================================

# Text file extensions (pure text, no embedded images by definition)
TEXT_EXTENSIONS = {
    ".txt", ".md", ".rst", ".csv", ".tsv",
    ".json", ".yaml", ".yml", ".xml", ".html", ".htm",
    ".rtf", ".log", ".ini", ".cfg", ".conf",
}

# Code file extensions
CODE_EXTENSIONS = {
    ".py", ".js", ".ts", ".jsx", ".tsx",
    ".java", ".kt", ".scala", ".groovy",
    ".c", ".cpp", ".cc", ".h", ".hpp", ".cs",
    ".go", ".rs", ".rb", ".php",
    ".swift", ".m", ".mm",
    ".sql", ".sh", ".bash", ".ps1", ".bat", ".cmd",
    ".r", ".R", ".jl", ".lua", ".pl", ".pm",
    ".hs", ".elm", ".ex", ".exs", ".erl",
    ".vue", ".svelte", ".astro",
    ".toml", ".dockerfile", ".makefile",
}

# Image file extensions
IMAGE_EXTENSIONS = {
    ".png", ".jpg", ".jpeg", ".gif", ".webp",
    ".bmp", ".tiff", ".tif", ".svg", ".ico",
}

# Video file extensions
VIDEO_EXTENSIONS = {
    ".mp4", ".mov", ".avi", ".mkv", ".webm",
    ".m4v", ".wmv", ".flv", ".mpeg", ".mpg",
}

# Document extensions that MAY contain embedded images (need inspection)
POTENTIALLY_MIXED_EXTENSIONS = {
    ".pdf", ".docx", ".doc", ".pptx", ".ppt",
    ".xlsx", ".xls", ".odt", ".odp", ".ods",
}

# Audio extensions (for completeness, routed separately)
AUDIO_EXTENSIONS = {
    ".mp3", ".wav", ".ogg", ".flac", ".m4a",
    ".aac", ".wma", ".aiff",
}


# =============================================================================
# CLASSIFIED FILE INFO
# =============================================================================

@dataclass
class ClassifiedFile:
    """
    A file with its classification and metadata.
    
    This is the output of file classification - contains everything
    needed for routing and context building.
    """
    # Stable identifier (e.g., "[FILE_1]")
    file_id: str
    
    # Original filename
    original_name: str
    
    # Classification
    file_type: FileType
    
    # File metadata
    extension: str
    size_bytes: int
    mime_type: Optional[str] = None
    
    # For MIXED_FILE: details about embedded content
    embedded_image_count: int = 0
    text_chars: int = 0
    page_count: int = 0
    
    # Path to file (for preprocessing)
    file_path: Optional[str] = None
    
    # Raw bytes (if loaded)
    file_bytes: Optional[bytes] = None
    
    # Extracted text content (populated by preprocessor)
    extracted_text: Optional[str] = None
    
    # Additional metadata
    metadata: Dict[str, Any] = field(default_factory=dict)
    
    def __str__(self) -> str:
        """Human-readable representation for file map."""
        if self.file_type == FileType.MIXED_FILE:
            return f"{self.file_id} {self.original_name} (MIXED_FILE: text + {self.embedded_image_count} images)"
        return f"{self.file_id} {self.original_name} ({self.file_type.value})"
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for JSON serialization."""
        return {
            "file_id": self.file_id,
            "original_name": self.original_name,
            "file_type": self.file_type.value,
            "extension": self.extension,
            "size_bytes": self.size_bytes,
            "mime_type": self.mime_type,
            "embedded_image_count": self.embedded_image_count,
            "text_chars": self.text_chars,
            "page_count": self.page_count,
            "has_extracted_text": self.extracted_text is not None,
        }


# =============================================================================
# CLASSIFICATION RESULT
# =============================================================================

@dataclass
class ClassificationResult:
    """
    Complete result of file classification.
    
    Contains:
    - All classified files with stable IDs
    - Canonical buckets (lists of files by type)
    - Modality flags
    """
    # All classified files in upload order
    classified_files: List[ClassifiedFile]
    
    # Canonical buckets (Spec §1.1)
    text_files: List[ClassifiedFile] = field(default_factory=list)
    code_files: List[ClassifiedFile] = field(default_factory=list)
    image_files: List[ClassifiedFile] = field(default_factory=list)
    video_files: List[ClassifiedFile] = field(default_factory=list)
    mixed_files: List[ClassifiedFile] = field(default_factory=list)
    
    # Modality flags (Spec §1.2)
    has_text: bool = False
    has_code: bool = False
    has_image: bool = False
    has_video: bool = False
    has_mixed: bool = False
    
    # Counts
    total_files: int = 0
    total_size_bytes: int = 0
    
    def to_dict(self) -> Dict[str, Any]:
        """Convert to dictionary for logging/serialization."""
        return {
            "total_files": self.total_files,
            "total_size_bytes": self.total_size_bytes,
            "flags": {
                "HAS_TEXT": self.has_text,
                "HAS_CODE": self.has_code,
                "HAS_IMAGE": self.has_image,
                "HAS_VIDEO": self.has_video,
                "HAS_MIXED": self.has_mixed,
            },
            "buckets": {
                "text_files": len(self.text_files),
                "code_files": len(self.code_files),
                "image_files": len(self.image_files),
                "video_files": len(self.video_files),
                "mixed_files": len(self.mixed_files),
            },
            "files": [f.to_dict() for f in self.classified_files],
        }


# =============================================================================
# IMAGE DETECTION IN DOCUMENTS
# =============================================================================

def count_images_in_pdf(
    pdf_path: Optional[str] = None,
    pdf_bytes: Optional[bytes] = None
) -> Tuple[int, int, int]:
    """
    Count images in a PDF file.
    
    Args:
        pdf_path: Path to PDF file
        pdf_bytes: Raw PDF bytes
    
    Returns:
        (image_count, text_chars, page_count)
    """
    try:
        # Try PyMuPDF first (most accurate)
        try:
            import fitz
            if pdf_bytes:
                doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            else:
                doc = fitz.open(pdf_path)
            
            try:
                page_count = len(doc)
                text_chars = 0
                image_count = 0
                
                for page in doc:
                    text_chars += len(page.get_text())
                    image_count += len(page.get_images(full=True))
                
                return image_count, text_chars, page_count
            finally:
                doc.close()
        except ImportError:
            pass
        
        # Fallback to pypdf
        try:
            from pypdf import PdfReader
            if pdf_bytes:
                reader = PdfReader(io.BytesIO(pdf_bytes))
            else:
                reader = PdfReader(pdf_path)
            
            page_count = len(reader.pages)
            text_chars = 0
            image_count = 0
            
            for page in reader.pages:
                text_chars += len(page.extract_text() or "")
                try:
                    if "/Resources" in page:
                        resources = page["/Resources"]
                        if "/XObject" in resources:
                            xobjects = resources["/XObject"]
                            if hasattr(xobjects, "get_object"):
                                xobjects = xobjects.get_object()
                            for obj_name in xobjects:
                                obj = xobjects[obj_name]
                                if hasattr(obj, "get_object"):
                                    obj = obj.get_object()
                                if obj.get("/Subtype") == "/Image":
                                    image_count += 1
                except Exception:
                    pass
            
            return image_count, text_chars, page_count
        except ImportError:
            pass
        
        logger.warning("No PDF library available for image detection")
        return 0, 0, 0
        
    except Exception as e:
        logger.error(f"Error counting images in PDF: {e}")
        return 0, 0, 0


def count_images_in_docx(
    docx_path: Optional[str] = None,
    docx_bytes: Optional[bytes] = None
) -> Tuple[int, int]:
    """
    Count images in a DOCX file.
    
    Args:
        docx_path: Path to DOCX file
        docx_bytes: Raw DOCX bytes
    
    Returns:
        (image_count, text_chars)
    """
    try:
        try:
            from docx import Document
            import zipfile
            
            # DOCX is a ZIP file - count image files in word/media/
            if docx_bytes:
                with zipfile.ZipFile(io.BytesIO(docx_bytes)) as zf:
                    image_count = sum(
                        1 for name in zf.namelist()
                        if name.startswith("word/media/") and 
                        any(name.lower().endswith(ext) for ext in [".png", ".jpg", ".jpeg", ".gif", ".bmp"])
                    )
                doc = Document(io.BytesIO(docx_bytes))
            else:
                with zipfile.ZipFile(docx_path) as zf:
                    image_count = sum(
                        1 for name in zf.namelist()
                        if name.startswith("word/media/") and 
                        any(name.lower().endswith(ext) for ext in [".png", ".jpg", ".jpeg", ".gif", ".bmp"])
                    )
                doc = Document(docx_path)
            
            # Count text characters
            text_chars = sum(len(p.text) for p in doc.paragraphs)
            
            return image_count, text_chars
            
        except ImportError:
            logger.warning("python-docx not available for DOCX image detection")
            return 0, 0
            
    except Exception as e:
        logger.warning(f"Error counting images in DOCX: {e}")
        return 0, 0


def count_images_in_pptx(
    pptx_path: Optional[str] = None,
    pptx_bytes: Optional[bytes] = None
) -> Tuple[int, int]:
    """
    Count images in a PPTX file.
    
    Args:
        pptx_path: Path to PPTX file
        pptx_bytes: Raw PPTX bytes
    
    Returns:
        (image_count, text_chars)
    """
    try:
        import zipfile
        
        # PPTX is a ZIP file - count image files in ppt/media/
        if pptx_bytes:
            zf = zipfile.ZipFile(io.BytesIO(pptx_bytes))
        else:
            zf = zipfile.ZipFile(pptx_path)
        
        with zf:
            image_count = sum(
                1 for name in zf.namelist()
                if name.startswith("ppt/media/") and 
                any(name.lower().endswith(ext) for ext in [".png", ".jpg", ".jpeg", ".gif", ".bmp"])
            )
        
        # PPTX always has images (slides), treat as mixed
        return image_count, 0
        
    except Exception as e:
        logger.warning(f"Error counting images in PPTX: {e}")
        return 0, 0


# =============================================================================
# MAIN CLASSIFICATION FUNCTION
# =============================================================================

def classify_single_file(
    filename: str,
    file_id: str,
    size_bytes: int = 0,
    mime_type: Optional[str] = None,
    file_path: Optional[str] = None,
    file_bytes: Optional[bytes] = None,
    pdf_image_count: Optional[int] = None,
    pdf_text_chars: Optional[int] = None,
    pdf_page_count: Optional[int] = None,
) -> ClassifiedFile:
    """
    Classify a single file into exactly one bucket.
    
    Args:
        filename: Original filename
        file_id: Stable identifier (e.g., "[FILE_1]")
        size_bytes: File size
        mime_type: MIME type hint
        file_path: Path to file (for inspection)
        file_bytes: Raw bytes (for inspection)
        pdf_image_count: Pre-computed PDF image count (from AttachmentInfo)
        pdf_text_chars: Pre-computed PDF text chars
        pdf_page_count: Pre-computed PDF page count
    
    Returns:
        ClassifiedFile with file_type set
    """
    ext = Path(filename).suffix.lower() if filename else ""
    
    # Default result
    result = ClassifiedFile(
        file_id=file_id,
        original_name=filename,
        file_type=FileType.TEXT_FILE,  # Default
        extension=ext,
        size_bytes=size_bytes,
        mime_type=mime_type,
        file_path=file_path,
        file_bytes=file_bytes,
    )
    
    # === CLASSIFICATION LOGIC ===
    
    # 1. VIDEO_FILE - check first (highest priority for routing)
    if ext in VIDEO_EXTENSIONS:
        result.file_type = FileType.VIDEO_FILE
        logger.debug(f"[file_classifier] {filename} → VIDEO_FILE (extension)")
        return result
    
    # 2. IMAGE_FILE - standalone images
    if ext in IMAGE_EXTENSIONS:
        result.file_type = FileType.IMAGE_FILE
        logger.debug(f"[file_classifier] {filename} → IMAGE_FILE (extension)")
        return result
    
    # 3. CODE_FILE - programming files
    if ext in CODE_EXTENSIONS:
        result.file_type = FileType.CODE_FILE
        logger.debug(f"[file_classifier] {filename} → CODE_FILE (extension)")
        return result
    
    # 4. POTENTIALLY_MIXED - documents that may contain images
    if ext in POTENTIALLY_MIXED_EXTENSIONS:
        # PDF handling
        if ext == ".pdf":
            image_count = pdf_image_count
            text_chars = pdf_text_chars or 0
            page_count = pdf_page_count or 0
            
            # If not pre-computed, try to analyze
            if image_count is None and (file_path or file_bytes):
                image_count, text_chars, page_count = count_images_in_pdf(file_path, file_bytes)
            
            result.embedded_image_count = image_count or 0
            result.text_chars = text_chars
            result.page_count = page_count
            
            if result.embedded_image_count > 0:
                result.file_type = FileType.MIXED_FILE
                logger.debug(f"[file_classifier] {filename} → MIXED_FILE (PDF with {result.embedded_image_count} images)")
            else:
                result.file_type = FileType.TEXT_FILE
                logger.debug(f"[file_classifier] {filename} → TEXT_FILE (PDF, no images)")
            return result
        
        # DOCX handling
        if ext in {".docx", ".doc"}:
            if file_path or file_bytes:
                image_count, text_chars = count_images_in_docx(file_path, file_bytes)
                result.embedded_image_count = image_count
                result.text_chars = text_chars
                
                if image_count > 0:
                    result.file_type = FileType.MIXED_FILE
                    logger.debug(f"[file_classifier] {filename} → MIXED_FILE (DOCX with {image_count} images)")
                else:
                    result.file_type = FileType.TEXT_FILE
                    logger.debug(f"[file_classifier] {filename} → TEXT_FILE (DOCX, no images)")
            else:
                # Can't inspect, assume TEXT_FILE
                result.file_type = FileType.TEXT_FILE
                logger.debug(f"[file_classifier] {filename} → TEXT_FILE (DOCX, assumed no images)")
            return result
        
        # PPTX handling - always has images (slides)
        if ext in {".pptx", ".ppt"}:
            if file_path or file_bytes:
                image_count, _ = count_images_in_pptx(file_path, file_bytes)
                result.embedded_image_count = max(image_count, 1)  # At least 1 for slides
            else:
                result.embedded_image_count = 1  # Assume has images
            result.file_type = FileType.MIXED_FILE
            logger.debug(f"[file_classifier] {filename} → MIXED_FILE (PPTX)")
            return result
        
        # XLSX/XLS - treat as text (no vision needed for spreadsheets)
        if ext in {".xlsx", ".xls"}:
            result.file_type = FileType.TEXT_FILE
            logger.debug(f"[file_classifier] {filename} → TEXT_FILE (spreadsheet)")
            return result
        
        # Other potentially mixed - can't inspect, assume text
        result.file_type = FileType.TEXT_FILE
        logger.debug(f"[file_classifier] {filename} → TEXT_FILE (document, assumed no images)")
        return result
    
    # 5. TEXT_FILE - known text extensions
    if ext in TEXT_EXTENSIONS:
        result.file_type = FileType.TEXT_FILE
        logger.debug(f"[file_classifier] {filename} → TEXT_FILE (extension)")
        return result
    
    # 6. AUDIO - not in our buckets, treat as text (will need transcription)
    if ext in AUDIO_EXTENSIONS:
        result.file_type = FileType.TEXT_FILE  # Will need audio transcription
        result.metadata["is_audio"] = True
        logger.debug(f"[file_classifier] {filename} → TEXT_FILE (audio, needs transcription)")
        return result
    
    # 7. Unknown extension - use MIME type or default to text
    if mime_type:
        if mime_type.startswith("video/"):
            result.file_type = FileType.VIDEO_FILE
        elif mime_type.startswith("image/"):
            result.file_type = FileType.IMAGE_FILE
        elif mime_type.startswith("text/") or mime_type == "application/json":
            result.file_type = FileType.TEXT_FILE
        else:
            result.file_type = FileType.TEXT_FILE
    else:
        result.file_type = FileType.TEXT_FILE
    
    logger.debug(f"[file_classifier] {filename} → {result.file_type.value} (fallback)")
    return result


def classify_attachments(
    attachments: List[Dict[str, Any]],
    base_path: Optional[str] = None,
) -> ClassificationResult:
    """
    Classify all attachments and build canonical buckets.
    
    Args:
        attachments: List of attachment dicts with keys:
            - filename: str (required)
            - mime_type: str (optional)
            - size_bytes: int (optional)
            - path: str (optional, for file inspection)
            - bytes: bytes (optional, for file inspection)
            - pdf_image_count: int (optional, pre-computed)
            - pdf_text_chars: int (optional)
            - pdf_page_count: int (optional)
        base_path: Base path for resolving relative file paths
    
    Returns:
        ClassificationResult with all files classified and buckets filled
    """
    classified_files: List[ClassifiedFile] = []
    
    # Classify each file
    for idx, att in enumerate(attachments):
        file_id = f"[FILE_{idx + 1}]"
        filename = att.get("filename", att.get("name", f"unknown_{idx}"))
        
        # Resolve file path
        file_path = att.get("path")
        if file_path and base_path and not os.path.isabs(file_path):
            file_path = os.path.join(base_path, file_path)
        
        classified = classify_single_file(
            filename=filename,
            file_id=file_id,
            size_bytes=att.get("size_bytes", att.get("size", 0)),
            mime_type=att.get("mime_type", att.get("type")),
            file_path=file_path,
            file_bytes=att.get("bytes", att.get("data")),
            pdf_image_count=att.get("pdf_image_count"),
            pdf_text_chars=att.get("pdf_text_chars"),
            pdf_page_count=att.get("pdf_page_count"),
        )
        
        classified_files.append(classified)
    
    # Build result with buckets
    result = ClassificationResult(classified_files=classified_files)
    
    # Fill buckets
    for cf in classified_files:
        if cf.file_type == FileType.TEXT_FILE:
            result.text_files.append(cf)
        elif cf.file_type == FileType.CODE_FILE:
            result.code_files.append(cf)
        elif cf.file_type == FileType.IMAGE_FILE:
            result.image_files.append(cf)
        elif cf.file_type == FileType.VIDEO_FILE:
            result.video_files.append(cf)
        elif cf.file_type == FileType.MIXED_FILE:
            result.mixed_files.append(cf)
    
    # Compute modality flags (Spec §1.2)
    result.has_text = len(result.text_files) > 0
    result.has_code = len(result.code_files) > 0
    result.has_image = len(result.image_files) > 0 or len(result.mixed_files) > 0
    result.has_video = len(result.video_files) > 0
    result.has_mixed = len(result.mixed_files) > 0
    
    # Counts
    result.total_files = len(classified_files)
    result.total_size_bytes = sum(cf.size_bytes for cf in classified_files)
    
    logger.info(
        f"[file_classifier] Classified {result.total_files} files: "
        f"text={len(result.text_files)}, code={len(result.code_files)}, "
        f"image={len(result.image_files)}, video={len(result.video_files)}, "
        f"mixed={len(result.mixed_files)}"
    )
    
    return result


# =============================================================================
# FILE MAP BUILDER (Spec §2)
# =============================================================================

def build_file_map(classified_files: List[ClassifiedFile]) -> str:
    """
    Build the stable file map string for model prompts.
    
    Format:
        FILE MAP:
        [FILE_1] 2025-12-10_bug_demo.mp4 (VIDEO_FILE)
        [FILE_2] main_router.py (CODE_FILE)
        [FILE_3] design_spec.pdf (MIXED_FILE: text + 5 images)
        ...
        When referring to files, always use [FILE_X].
    
    Args:
        classified_files: List of ClassifiedFile objects
    
    Returns:
        Formatted file map string
    """
    if not classified_files:
        return "FILE MAP:\n(no files attached)\n"
    
    lines = ["FILE MAP:"]
    for cf in classified_files:
        lines.append(str(cf))
    lines.append("When referring to files, always use [FILE_X].")
    
    return "\n".join(lines)


def build_file_map_for_task(
    classified_files: List[ClassifiedFile],
    target_file_ids: List[str],
) -> str:
    """
    Build file map for a specific task (subset of files).
    
    Args:
        classified_files: All classified files
        target_file_ids: List of file IDs for this task (e.g., ["[FILE_1]", "[FILE_3]"])
    
    Returns:
        Formatted file map string with only relevant files
    """
    relevant = [cf for cf in classified_files if cf.file_id in target_file_ids]
    return build_file_map(relevant)


# =============================================================================
# HELPER FUNCTIONS
# =============================================================================

def get_file_by_id(
    classified_files: List[ClassifiedFile],
    file_id: str,
) -> Optional[ClassifiedFile]:
    """Get a classified file by its stable ID."""
    for cf in classified_files:
        if cf.file_id == file_id:
            return cf
    return None


def get_files_by_type(
    classified_files: List[ClassifiedFile],
    file_type: FileType,
) -> List[ClassifiedFile]:
    """Get all files of a specific type."""
    return [cf for cf in classified_files if cf.file_type == file_type]


def has_any_media(result: ClassificationResult) -> bool:
    """Check if classification has any media (images or video)."""
    return result.has_image or result.has_video


def has_vision_content(result: ClassificationResult) -> bool:
    """Check if classification requires vision model (images, video, or mixed)."""
    return result.has_image or result.has_video or result.has_mixed


# =============================================================================
# CONVERSION FROM EXISTING ATTACHMENTINFO
# =============================================================================

def classify_from_attachment_info(
    attachment_infos: List[Any],
) -> ClassificationResult:
    """
    Convert from existing AttachmentInfo objects to classification result.
    
    This bridges the old AttachmentInfo schema to the new file classification system.
    
    Args:
        attachment_infos: List of AttachmentInfo objects (from app.llm.schemas)
    
    Returns:
        ClassificationResult
    """
    attachments = []
    
    for att in attachment_infos:
        attachments.append({
            "filename": getattr(att, "filename", "unknown"),
            "mime_type": getattr(att, "mime_type", None),
            "size_bytes": getattr(att, "size_bytes", 0),
            "pdf_image_count": getattr(att, "pdf_image_count", None),
            "pdf_text_chars": getattr(att, "pdf_text_chars", None),
            "pdf_page_count": getattr(att, "pdf_page_count", None),
        })
    
    return classify_attachments(attachments)


# =============================================================================
# EXPORTS
# =============================================================================

__all__ = [
    # Enums
    "FileType",
    
    # Data classes
    "ClassifiedFile",
    "ClassificationResult",
    
    # Main functions
    "classify_single_file",
    "classify_attachments",
    "classify_from_attachment_info",
    
    # File map
    "build_file_map",
    "build_file_map_for_task",
    
    # Helpers
    "get_file_by_id",
    "get_files_by_type",
    "has_any_media",
    "has_vision_content",
    
    # Image detection
    "count_images_in_pdf",
    "count_images_in_docx",
    "count_images_in_pptx",
    
    # Extension sets (for external use)
    "TEXT_EXTENSIONS",
    "CODE_EXTENSIONS",
    "IMAGE_EXTENSIONS",
    "VIDEO_EXTENSIONS",
    "POTENTIALLY_MIXED_EXTENSIONS",
]